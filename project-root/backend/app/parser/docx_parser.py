import docx
import re
from typing import BinaryIO, List, Dict, Any
from .base import BaseParser, ParsedDocument
from uuid import uuid4

class DOCXParser(BaseParser):
    def supports(self, file_type: str) -> bool:
        return file_type.lower() == 'docx'
    
    def parse(self, file_stream: BinaryIO, file_name: str) -> ParsedDocument:
        doc = docx.Document(file_stream)
        
        full_text = []
        sections = []
        current_section = None
        
        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue
            
            style = para.style.name if para.style else "Normal"
            level = self._detect_heading_level(style, text)
            
            if level > 0:
                section = {
                    "number": self._extract_section_number(text),
                    "title": self._clean_title(text),
                    "level": level,
                    "page": None,  # DOCX не даёт номер страницы просто так
                    "text": []
                }
                sections.append(section)
                current_section = section
            else:
                full_text.append(text)
                if current_section:
                    current_section["text"].append(text)
        
        for section in sections:
            section["text"] = "\n".join(section["text"])
        
        content = "\n".join(full_text)
        
        return ParsedDocument(
            document_id=uuid4(),
            file_name=file_name,
            file_type="docx",
            content=content,
            sections=sections,
            metadata={
                "paragraph_count": len(doc.paragraphs),
                "title": doc.core_properties.title or "",
                "author": doc.core_properties.author or ""
            },
            entities=self._extract_entities(content)
        )
    
    def _detect_heading_level(self, style_name: str, text: str) -> int:
        """Определяет уровень по стилю Word"""
        style_lower = style_name.lower()
        
        if 'heading 1' in style_lower or style_lower == 'title':
            return 1
        elif 'heading 2' in style_lower:
            return 2
        elif 'heading 3' in style_lower:
            return 3
        
        # Fallback по числовому паттерну
        if re.match(r'^\s*\d+\s+', text) and len(text) < 100:
            return 1
        if re.match(r'^\s*\d+\.\d+\s+', text) and len(text) < 100:
            return 2
        
        return 0
    
    def _extract_section_number(self, text: str) -> str:
        match = re.match(r'^\s*(\d+(?:\.\d+)*)\s+', text)
        return match.group(1) if match else ""
    
    def _clean_title(self, text: str) -> str:
        return re.sub(r'^\s*\d+(?:\.\d+)*\s+', '', text).strip()
    
    def _extract_entities(self, text: str) -> List[Dict[str, Any]]:
        # Та же логика, что и в PDFParser
        entities = []
        for match in re.finditer(r'ГОСТ\s+(?:Р\s+)?\d+(?:-\d+)+', text):
            entities.append({
                "type": "standard",
                "value": match.group(),
                "position": match.span()
            })
        return entities